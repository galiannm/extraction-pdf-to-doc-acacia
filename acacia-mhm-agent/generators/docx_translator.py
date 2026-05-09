"""
Translate Acacia_MHM_FR.docx → Acacia_MHM_EN.docx by:
1. Copying the FR docx verbatim (all styling preserved)
2. Translating all body text in-place
3. Re-applying English MHM term highlighting on body paragraphs
"""
import json
import os
import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from dotenv import load_dotenv
from google import genai
from google.genai import types as gtypes
from rich.console import Console

from config import Fonts
from templates.acacia_styles import _add_highlighted_runs

load_dotenv()
console = Console()

BATCH_SIZE = 10

_PROMPT = """You are a native English-speaking academic expert in early childhood education and pedagogy.
Your task is to produce a high-quality English version of each French text below.

IMPORTANT:
- This is NOT a literal translation
- This is NOT a simplification
- You must preserve ALL key ideas, concepts, and pedagogical intentions

TRANSLATION APPROACH:
- Rewrite the content in natural, fluent, professional English
- Ensure it reads as if it was originally written by a native English-speaking education expert
- Use appropriate educational terminology used in international schools
- Maintain structure, sections, and pedagogical clarity

STYLE REQUIREMENTS:
- Academic but accessible tone
- Clear, precise, and professional language
- Natural flow (no "translated feeling")
- Consistent vocabulary throughout

CRITICAL GOAL:
The final English version must give the impression that:
→ it was originally written in English
→ by a native expert in early childhood education
NOT translated from French

RULES:
- Return ONLY valid JSON: {"translations": ["...", "...", ...]}
- The array must contain EXACTLY the same number of items as the input array
- Preserve bullet symbols (•, -) at the start of lines exactly as they appear
- Preserve tab characters (\\t) where they appear in the input
- Preserve newlines (\\n) in the output where they appear in the input
- Do NOT translate: codes like N1, G2, E4, P1–P5, the letter X, page numbers, or URLs
- If an item is already English, a code, or a symbol — return it unchanged
- Translate "Semaine" → "Week", "Période" → "Period", "Petite Section" → "Kindergarten 1", "Moyenne Section" → "Kindergarten 2"
- Translate "Jour" → "Day", "CHAQUE JOUR" → "DAILY", "JOUR 1" → "DAY 1", "JOUR 2" → "DAY 2", etc.
- Translate "Ce qu'il faut savoir" → "What you need to know"
- Translate "Différenciation" → "Differentiation", "Différenciation et évolution de l'activité" → "Differentiation and activity progression"
"""


def translate_docx(
    fr_path: Path,
    en_path: Path,
    title: str,
    subtitle: str,
    fr_title: str = "Guide des séances MHM",
    fr_subtitle: str = "Petite Section / Moyenne Section — Acacia International Pre-school",
    translation_prompt: str | None = None,
) -> Path:
    """Copy FR docx → translate all body text in-place → save as EN docx."""
    console.print(f"[cyan]Translating docx:[/cyan] {fr_path.name} → {en_path.name}")
    shutil.copy(str(fr_path), str(en_path))
    doc = Document(str(en_path))
    client = _make_client()
    _active_prompt = translation_prompt if translation_prompt is not None else _PROMPT

    items = _collect_items(doc)
    console.print(f"  {len(items)} text segments collected")

    texts = [text for _, text in items]
    translated_texts = _batch_translate(client, texts, _active_prompt)

    for (para, _), new_text in zip(items, translated_texts):
        if new_text and new_text.strip():
            _apply_to_para(para, new_text)

    # Retry items where the translation returned the original text unchanged
    skipped = [(para, orig) for (para, orig), trans in zip(items, translated_texts)
               if trans == orig and _should_translate(orig)]
    if skipped:
        console.print(f"  [yellow]Retrying {len(skipped)} unchanged items in batches of 5\u2026[/yellow]")
        skip_paras  = [p for p, _ in skipped]
        skip_texts  = [t for _, t in skipped]
        skip_results = []
        for i in range(0, len(skip_texts), 5):
            sub = skip_texts[i:i+5]
            skip_results.extend(_call_gemini(client, sub, _active_prompt))
        for para, new_text in zip(skip_paras, skip_results):
            if new_text and new_text.strip():
                _apply_to_para(para, new_text)

    # Final passes: detect paragraphs still in French and retranslate (run twice to catch stragglers)
    _retranslate_french_paragraphs(doc, client, _active_prompt)
    _retranslate_french_paragraphs(doc, client, _active_prompt)

    _patch_cover(doc, fr_title, title, fr_subtitle, subtitle)
    _fix_section_labels(doc)

    doc.save(str(en_path))
    console.print(f"  [green]Saved →[/green] {en_path}")
    return en_path


# ── Collection ────────────────────────────────────────────────────────────────

def _collect_items(doc: Document) -> list[tuple]:
    items = []
    seen = set()

    def add(para):
        pid = id(para)
        if pid in seen:
            return
        seen.add(pid)
        text = para.text
        if _should_translate(text):
            items.append((para, text))

    for para in doc.paragraphs:
        add(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    add(para)

    return items


def _should_translate(text: str) -> bool:
    t = text.strip()
    if not t or len(t) <= 2:
        return False
    if re.fullmatch(r'[NnGgEe]\d+', t):   # competence codes
        return False
    if re.fullmatch(r'\d+', t):            # page numbers
        return False
    if t.lower().startswith('http'):
        return False
    if re.fullmatch(r'[Xx]', t):           # X marks in tables
        return False
    if re.fullmatch(r'[Ss]\d+', t):        # week codes S1–S5
        return False
    # Skip period/week banner cells like "Période 1\nS1" — keep structural codes intact
    if re.search(r'\b[Ss]\d+\s*$', t) and re.search(r'[Pp]ériode', t):
        return False
    return True


# ── Translation ───────────────────────────────────────────────────────────────

def _batch_translate(client, texts: list[str], prompt_template: str = _PROMPT) -> list[str]:
    results = []
    total = len(texts)
    for i in range(0, total, BATCH_SIZE):
        batch = texts[i : i + BATCH_SIZE]
        n = len(results) + len(batch)
        console.print(f"  Batch {i // BATCH_SIZE + 1}: items {i+1}–{n}/{total}…")
        translated = _call_gemini(client, batch, prompt_template)

        # Tier 2: if full batch fell back, retry in sub-batches of 10
        if translated == batch and len(batch) > 1:
            console.print("  [yellow]Retrying in sub-batches of 10…[/yellow]")
            translated = []
            for j in range(0, len(batch), 10):
                sub = batch[j : j + 10]
                sub_result = _call_gemini(client, sub, prompt_template)
                # Tier 3: if sub-batch also fell back, translate item by item
                if sub_result == sub and len(sub) > 1:
                    sub_result = []
                    for item in sub:
                        sub_result.extend(_call_gemini(client, [item], prompt_template))
                translated.extend(sub_result)

        results.extend(translated)
    return results


def _call_gemini(client, texts: list[str], prompt_template: str = _PROMPT) -> list[str]:
    payload = json.dumps({"texts": texts}, ensure_ascii=False)
    prompt = f"{prompt_template}\n\nTexts to translate:\n{payload}"

    for attempt in range(1, 4):
        try:
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=prompt,
                config=gtypes.GenerateContentConfig(
                    temperature=0,
                    thinking_config=gtypes.ThinkingConfig(thinking_budget=0),
                ),
            )
            raw = response.text.strip()
            if raw.startswith("```"):
                raw = raw.split("\n", 1)[1].rsplit("```", 1)[0]
            data = json.loads(_clean_json(raw))
            result = data.get("translations", [])
            if len(result) == len(texts):
                return result
            console.print(f"  [yellow]Length mismatch ({len(result)} vs {len(texts)}), retry {attempt}…[/yellow]")
        except Exception as e:
            console.print(f"  [red]Attempt {attempt} failed:[/red] {e}")

    console.print("  [yellow]Returning originals as fallback[/yellow]")
    return texts


# ── Application ───────────────────────────────────────────────────────────────

def _apply_to_para(para, new_text: str) -> None:
    runs = para.runs
    if not runs:
        return
    # Always preserve original run styling: put translated text in first run, clear the rest.
    # This faithfully copies the FR doc styling (font, color, size) into the EN doc.
    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


_SECTION_LABEL_MAP = {
    "ce qu'il faut savoir":                          "What you need to know",
    "différenciation":                               "Differentiation",
    "différenciation et évolution de l'activité":    "Differentiation and activity progression",
    "déroulement":                                   "How it works",
    "matériel":                                      "Materials",
    "matériel :":                                    "Materials:",
    "objectif":                                      "Objective",
    "objectifs":                                     "Objectives",
    "objectif :":                                    "Objective:",
    "objectifs :":                                   "Objectives:",
    "organisation de classe":                        "Class organisation",
    "organisation de classe :":                      "Class organisation:",
    "ressources":                                    "Resources",
    "ressources :":                                  "Resources:",
    "matériel de classe":                            "Class materials",
    "remarques":                                     "Notes",
    "ce qu'il faut savoir":                          "What you need to know",
    "1er temps de travail de la journee":             "1st work session of the day",
    "1er temps de travail de la journee":             "1st work session of the day",
    "2e temps de travail de la journee":              "2nd work session of the day",
    "2eme temps de travail de la journee":            "2nd work session of the day",
    "2ème temps de travail de la journée":            "2nd work session of the day",
    "2e temps de travail de la journée":              "2nd work session of the day",
    "1er temps de travail de la journée":             "1st work session of the day",
}



def _still_french(text: str) -> bool:
    """Heuristic: text is likely still in French (not translated)."""
    if len(text.strip()) < 20:
        return False
    t = text.lower()
    # French-specific word patterns and accented characters (incl. curly apostrophe)
    indicators = [
        ' de ', ' du ', ' des ', ' les ', ' une ', ' un ', ' en ',
        ' et ', ' dans ', ' est ', ' le ', ' la ', ' au ',
        "l’", "d’", "qu’", "l'", "d'", "qu'",
        'é', 'è', 'à', 'â', 'ê', 'ï', 'ô', 'û', 'ç',
    ]
    count = sum(1 for ind in indicators if ind in t)
    return count >= 2


def _retranslate_french_paragraphs(doc, client, prompt: str) -> None:
    """Scan all paragraphs for ones still in French and retranslate them."""
    french_paras = []
    seen = set()

    def _scan(para):
        pid = id(para)
        if pid in seen:
            return
        seen.add(pid)
        if _still_french(para.text) and _should_translate(para.text):
            french_paras.append(para)

    for para in doc.paragraphs:
        _scan(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _scan(para)

    if not french_paras:
        return

    from rich.console import Console as _Console
    _console = _Console()
    _console.print(f"  [yellow]Final pass: re-translating {len(french_paras)} French paragraphs…[/yellow]")

    # Translate one at a time — guarantees no length mismatches or silent skips
    for para in french_paras:
        text = para.text
        result = _call_gemini(client, [text], prompt)
        new_text = result[0] if result else text
        if new_text and new_text.strip() and new_text != text:
            _apply_to_para(para, new_text)


def _fix_section_labels(doc) -> None:
    """Post-pass: guarantee known section labels and badge texts are in English."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t_orig = para.text.strip()
                    t = t_orig.lower().replace('\u2019', "'").replace('\u2018', "'")
                    en = _SECTION_LABEL_MAP.get(t)
                    if not en:
                        t_up = t_orig.upper()
                        if t_up == 'CHAQUE JOUR':
                            en = 'DAILY'
                        elif re.fullmatch(r'JOUR\s+\d+', t_up):
                            en = 'DAY ' + re.search(r'\d+', t_up).group()
                    if en and para.runs:
                        para.runs[0].text = en
                        for run in para.runs[1:]:
                            run.text = ""


def _patch_cover(doc, fr_title, en_title, fr_subtitle, en_subtitle) -> None:
    for para in doc.paragraphs:
        t = para.text.strip()
        if t == fr_title:
            _simple_replace(para, en_title)
        elif t == fr_subtitle:
            _simple_replace(para, en_subtitle)


def _simple_replace(para, new_text: str) -> None:
    runs = para.runs
    if not runs:
        return
    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


# ── Helpers ───────────────────────────────────────────────────────────────────

def _clean_json(raw: str) -> str:
    # Normalise non-breaking and narrow no-break spaces before parsing
    raw = raw.replace(' ', ' ').replace(' ', ' ')
    result = []
    in_string = False
    escaped = False
    _ESCAPE = {'\n': '\\n', '\r': '\\r', '\t': '\\t'}
    for ch in raw:
        if escaped:
            result.append(ch)
            escaped = False
        elif ch == '\\' and in_string:
            result.append(ch)
            escaped = True
        elif ch == '"':
            result.append(ch)
            in_string = not in_string
        elif in_string and ord(ch) < 0x20:
            result.append(_ESCAPE.get(ch, ''))
        else:
            result.append(ch)
    return ''.join(result)


def _make_client():
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise EnvironmentError("GEMINI_API_KEY not set in .env")
    return genai.Client(
        api_key=api_key,
        http_options=gtypes.HttpOptions(timeout=120_000),
    )
