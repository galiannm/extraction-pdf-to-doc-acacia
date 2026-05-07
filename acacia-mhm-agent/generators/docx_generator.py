"""
Builds an Acacia-branded .docx from a list of extracted (or translated) page dicts.
The same generator is used for both the FR and EN document.
"""
import re
from pathlib import Path
from docx import Document
from rich.console import Console

import templates.acacia_styles as _acacia_styles
from templates.acacia_styles import (
    setup_document,
    add_header,
    add_footer,
    add_cover_page,
    add_period_banner,
    add_activity_title_bar,
    add_subtitle_badge,
    add_heading,
    add_section_label,
    add_info_box,
    add_section_box,
    add_obj_mat_table,
    add_materiel_table,
    add_prog_banner,
    add_image_table,
    add_paragraph,
    add_indented_paragraph,
    add_toc_table,
    add_table,
    add_image,
    add_image_placeholder,
)

console = Console()


def build_document(
    pages: list[dict],
    output_path: Path,
    title: str,
    subtitle: str = "",
    lang: str = "fr-FR",
) -> Path:
    doc = Document()
    setup_document(doc, lang)
    add_header(doc)
    add_footer(doc)
    add_cover_page(doc, title, subtitle)

    console.print(f"[cyan]Building document:[/cyan] {output_path.name}")
    pages = _inject_toc_table(pages)
    pages = _preprocess_pages(pages)
    total_blocks = sum(len(p["blocks"]) for p in pages)
    console.print(f"  {len(pages)} pages, {total_blocks} blocks")

    prev_was_bullet = False
    for page in pages:
        blocks = _mark_qr_labels(page["blocks"])
        for block in blocks:
            prev_was_bullet = _render_block(doc, block, lang, prev_was_bullet)
        # Page break after the cover/objectives page so Programmation starts fresh
        if page.get("page_num") == 1:
            doc.add_page_break()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    console.print(f"  [green]Saved →[/green] {output_path}")
    return output_path


def build_faithful_document(
    pages: list[dict],
    output_path: Path,
    title: str,
    subtitle: str = "",
    lang: str = "fr-FR",
) -> Path:
    """Build a Word doc that mimics the source PDF — standard fonts, no Acacia branding."""
    _acacia_styles._set_mode(faithful=True)
    try:
        result = build_document(pages, output_path, title, subtitle, lang)
    finally:
        _acacia_styles._set_mode(faithful=False)
    return result


# ── Page pre-processing ───────────────────────────────────────────────────────

def _is_differenciation(text: str) -> bool:
    t = text.lower().lstrip('💡 ')
    return t.startswith('différenciation') or t.startswith('differenciation')

# Block types that mark the start of a new section — stop collecting box content
_BOX_STOP_TYPES = {'heading', 'activity_title', 'subtitle_badge', 'section_box',
                   'toc_table', 'image_placeholder'}


def _group_section_boxes(blocks: list[dict]) -> list[dict]:
    """
    Group "Ce qu'il faut savoir" and "Différenciation" headings with their
    following content blocks into a single section_box block.
    """
    result = []
    i = 0
    while i < len(blocks):
        block = blocks[i]
        if block.get('type') != 'heading':
            result.append(block)
            i += 1
            continue

        text = (block.get('text') or '').strip()

        if _is_savoir_box(text):
            border = 'blue'
        elif _is_differenciation(text):
            border = 'orange'
        else:
            result.append(block)
            i += 1
            continue

        # Collect content blocks until the next section boundary
        content: list[dict] = []
        j = i + 1
        while j < len(blocks):
            b = blocks[j]
            if b.get('type') in _BOX_STOP_TYPES:
                break
            content.append(b)
            j += 1

        result.append({
            'type': 'section_box',
            'label': text,
            'border_color': border,
            'content_blocks': content,
        })
        i = j  # skip consumed content blocks

    return result


def _heading_norm(text: str) -> str:
    return re.sub(r'\*\*', '', text or '').lower().replace('’', "'").replace('‘', "'").strip()

def _is_objectif_heading(text: str) -> bool:
    return _heading_norm(text).startswith('objectif')

def _is_materiel_heading(text: str) -> bool:
    t = _heading_norm(text)
    return t.startswith('matériel') or t.startswith('materiel')

def _is_weekly_materiel(blocks: list, pos: int) -> bool:
    """True if the heading at pos is the weekly MATÉRIEL section (followed by PS/MS headings)."""
    for b in blocks[pos+1:pos+4]:
        if b.get('type') != 'heading':
            continue
        t = _heading_norm(b.get('text', ''))
        if t in ('ps', 'ms') or t.startswith('ps ') or t.startswith('ms ') or t.startswith('ps—') or t.startswith('ms—'):
            return True
    return False

_MAT_STOP_TYPES = {'activity_title', 'subtitle_badge', 'section_box', 'toc_table'}


def _group_obj_mat(blocks: list[dict]) -> list[dict]:
    """
    Group adjacent Objectif → (optional section_box) → Matériel into a single
    obj_mat_table block. The section_box (if any) is emitted before the table.
    """
    result = []
    i = 0
    while i < len(blocks):
        b = blocks[i]
        if b.get('type') != 'heading' or not _is_objectif_heading(b.get('text', '')):
            result.append(b)
            i += 1
            continue

        obj_label = b.get('text', '')
        obj_content: list[dict] = []
        interstitials: list[dict] = []
        j = i + 1

        # Collect objectif content until heading or section_box
        while j < len(blocks):
            nb = blocks[j]
            if nb.get('type') == 'heading':
                break
            if nb.get('type') == 'section_box':
                interstitials.append(nb)
                j += 1
                break
            obj_content.append(nb)
            j += 1

        # Skip any further section_boxes between objectif and matériel
        while j < len(blocks) and blocks[j].get('type') == 'section_box':
            interstitials.append(blocks[j])
            j += 1

        # Check if the next heading is Matériel (and NOT the weekly MATÉRIEL)
        if (j < len(blocks) and
                blocks[j].get('type') == 'heading' and
                _is_materiel_heading(blocks[j].get('text', '')) and
                not _is_weekly_materiel(blocks, j)):

            mat_label = blocks[j].get('text', '')
            mat_content: list[dict] = []
            k = j + 1
            while k < len(blocks):
                mb = blocks[k]
                if mb.get('type') in {'heading'} | _MAT_STOP_TYPES:
                    break
                mat_content.append(mb)
                k += 1

            result.extend(interstitials)
            result.append({
                'type': 'obj_mat_table',
                'obj_label': obj_label,
                'obj_blocks': obj_content,
                'mat_label': mat_label,
                'mat_blocks': mat_content,
            })
            i = k
            continue

        # No adjacent Matériel found — emit as-is
        result.append(b)
        result.extend(obj_content)
        result.extend(interstitials)
        i = j

    return result


def _group_materiel_table(blocks: list[dict]) -> list[dict]:
    """
    Detect the weekly MATÉRIEL overview section (MATÉRIEL heading followed by PS/MS
    headings) and convert it into a structured materiel_table block.
    """
    result = []
    i = 0
    while i < len(blocks):
        b = blocks[i]
        if (b.get('type') != 'heading' or
                not _is_materiel_heading(b.get('text', '')) or
                not _is_weekly_materiel(blocks, i)):
            result.append(b)
            i += 1
            continue

        # Collect all content until the next major boundary
        sections: dict[str, list] = {}
        current_key: str | None = None
        j = i + 1

        while j < len(blocks):
            nb = blocks[j]
            if nb.get('type') in _MAT_STOP_TYPES | {'activity_title', 'subtitle_badge'}:
                break
            if nb.get('type') == 'heading':
                t = _heading_norm(nb.get('text', ''))
                # Split-mode: "ps — ressources", "ms — matériel de classe", etc.
                if t.startswith('ps') and 'ressource' in t:
                    current_key = 'ps_ressources'
                elif t.startswith('ms') and 'ressource' in t:
                    current_key = 'ms_ressources'
                elif t.startswith('ps') and ('matériel' in t or 'materiel' in t or 'classe' in t):
                    current_key = 'ps_classe'
                elif t.startswith('ms') and ('matériel' in t or 'materiel' in t or 'classe' in t):
                    current_key = 'ms_classe'
                # Legacy: standalone PS/MS/Ressources headings
                elif t in ('ps', 'ms', 'ressource', 'ressources'):
                    current_key = t
                elif t.startswith('matériel de') or t.startswith('materiel de'):
                    current_key = 'mat_classe'
                else:
                    break
                sections.setdefault(current_key, [])
            elif nb.get('type') == 'paragraph' and current_key:
                sections[current_key].append(nb)
            j += 1

        result.append({
            'type': 'materiel_table',
            'ps_ressources':     sections.get('ps_ressources', []),
            'ms_ressources':     sections.get('ms_ressources', []),
            'ps_classe':         sections.get('ps_classe', []),
            'ms_classe':         sections.get('ms_classe', []),
            # Legacy fallback for pages not yet re-extracted
            'ressources_blocks': sections.get('ressource', sections.get('ressources', [])),
            'classe_blocks':     sections.get('mat_classe', []),
        })
        i = j

    return result


def _group_image_tables(blocks: list[dict]) -> list[dict]:
    """
    Group consecutive caption and image_placeholder blocks into a single
    image_table block rendered as a 1-row multi-column table.
    Skips page-number captions and very short/empty captions.
    """
    _IMAGE_TYPES = {'caption', 'image_placeholder'}

    def _caption_text(b: dict) -> str:
        raw = (b.get('text') or b.get('description') or '').strip()
        # Strip leading ► / ▸ arrows
        raw = re.sub(r'^[►▸]\s*', '', raw)
        # Skip pure page numbers
        if re.fullmatch(r'\d{1,3}', raw):
            return ''
        return raw

    result = []
    i = 0
    while i < len(blocks):
        b = blocks[i]
        if b.get('type') not in _IMAGE_TYPES:
            result.append(b)
            i += 1
            continue

        # Collect consecutive image/caption blocks
        descriptions = []
        j = i
        while j < len(blocks) and blocks[j].get('type') in _IMAGE_TYPES:
            text = _caption_text(blocks[j])
            if text and len(text) > 4:
                descriptions.append(text)
            j += 1

        if descriptions:
            result.append({'type': 'image_table', 'descriptions': descriptions})
        i = j

    return result


def _preprocess_pages(pages: list[dict]) -> list[dict]:
    result = []
    for page in pages:
        blocks = _merge_programmation_blocks(page.get("blocks", []))
        blocks = _group_section_boxes(blocks)
        blocks = _group_obj_mat(blocks)
        blocks = _group_materiel_table(blocks)
        blocks = _group_image_tables(blocks)
        new_page = dict(page)
        new_page["blocks"] = blocks
        result.append(new_page)
    return result


def _val_to_str(val) -> str:
    if val is None:
        return ""
    if isinstance(val, list):
        return ", ".join(str(v) for v in val if str(v).strip())
    return str(val)


def _content_blocks_to_cell(blocks: list[dict]) -> str:
    parts = []
    for b in blocks:
        btype = b.get("type", "")
        if btype == "paragraph":
            text = (_val_to_str(b.get("text"))).strip()
            if text:
                parts.append(text)
        elif btype == "heading":
            text = (_val_to_str(b.get("text"))).strip()
            if text:
                parts.append(f"**{text}**")
        elif btype == "subtitle_badge":
            badge = _val_to_str(b.get("badge_text", "")).strip().strip("[]'\"")
            title = _val_to_str(b.get("title_text", "")).strip()
            if badge and badge not in ("", "''"):
                parts.append(f"**{badge}** {title}".strip())
            elif title:
                parts.append(title)
    return "\n".join(p for p in parts if p.strip())


def _merge_programmation_blocks(blocks: list[dict]) -> list[dict]:
    """
    Detect the Programmation overview pattern (repeated activity_title blocks
    for the same week structure) and convert to a single table block.
    """
    activity_titles = [b for b in blocks if b.get("type") == "activity_title"]
    if len(activity_titles) < 6:
        return blocks

    # Must have at least 2 different activity types, each appearing multiple times
    type_counts: dict[str, int] = {}
    for b in activity_titles:
        t = b.get("activity_type", "")
        type_counts[t] = type_counts.get(t, 0) + 1
    if len(type_counts) < 2 or max(type_counts.values()) < 2:
        return blocks

    # Group blocks by (week_num, activity_type)
    groups: dict[tuple, list] = {}
    non_activity: list[dict] = []
    trailing: list[dict] = []
    current_title: dict | None = None
    current_content: list[dict] = []
    seen_first = False

    for block in blocks:
        if block.get("type") == "activity_title":
            if current_title is not None:
                wm = re.search(r"\d+", str(current_title.get("week", "")))
                wnum = int(wm.group()) if wm else 0
                key = (wnum, current_title.get("activity_type", ""))
                groups[key] = current_content
            else:
                non_activity = list(current_content)
            current_title = block
            current_content = []
            seen_first = True
        elif seen_first:
            current_content.append(block)
        else:
            non_activity.append(block)

    if current_title is not None:
        wm = re.search(r"\d+", str(current_title.get("week", "")))
        wnum = int(wm.group()) if wm else 0
        key = (wnum, current_title.get("activity_type", ""))
        groups[key] = current_content

    week_nums = sorted(set(k[0] for k in groups))
    _ATYPE_ORDER  = ["ritualisee", "guidee", "autonomie"]
    _ATYPE_LABELS = {
        "ritualisee": "Activités ritualisées",
        "guidee":     "Apprentissages guidés",
        "autonomie":  "Autonomie et semi-autonomie",
    }
    _ATYPE_COLORS = {"ritualisee": "yellow", "guidee": "blue", "autonomie": "green"}

    present = [a for a in _ATYPE_ORDER if any(a == k[1] for k in groups)]
    header = [""] + [_ATYPE_LABELS[a] for a in present]
    rows = [header]
    for wnum in week_nums:
        row = [f"S{wnum}"] + [_content_blocks_to_cell(groups.get((wnum, a), [])) for a in present]
        rows.append(row)

    col_colors = ["none"] + [_ATYPE_COLORS.get(a, "none") for a in present]

    # Extract periode_num from the first activity_title block in the original blocks
    periode_num = next(
        (b.get("periode_num", 1) for b in blocks if b.get("type") == "activity_title"),
        1,
    )

    overview_title = {
        "type": "prog_banner",
        "periode_num": periode_num,
        "label": "Programmation",
        "classes": ["PS", "MS"],
    }
    table_block = {"type": "table", "data": rows, "col_colors": col_colors, "compact": True}
    return non_activity + [overview_title, table_block] + trailing


# ── TOC pre-processing ────────────────────────────────────────────────────────

_TAB_PAGE_RE  = re.compile(r'\t(\d+)\.?\s*$')
_PERIOD_RE    = re.compile(r'^(PÉRIODE|PERIOD)\s*\d+', re.IGNORECASE)
_WEEK_RE      = re.compile(r'^(Semaine|Semame|Week)\s+', re.IGNORECASE)
_BULLET_STRIP = re.compile(r'^[•►▸\-]\s*')


def _parse_toc_line(line: str) -> dict | None:
    """Parse one TOC text line → entry dict or None."""
    line = line.strip()
    if not line:
        return None
    clean = _BULLET_STRIP.sub('', line)

    if _PERIOD_RE.match(clean):
        return {'level': 'period', 'title': clean, 'page': None}

    if _WEEK_RE.match(clean):
        m = _TAB_PAGE_RE.search(clean)
        if m:
            return {'level': 'item', 'title': clean[:m.start()].strip().rstrip('.'), 'page': int(m.group(1))}
        return {'level': 'week', 'title': clean, 'page': None}

    m = _TAB_PAGE_RE.search(clean)
    if m:
        title = clean[:m.start()].strip().rstrip('.')
        # Remove trailing dots/ellipsis from OCR artefacts
        title = re.sub(r'\.{2,}$', '', title).strip()
        return {'level': 'item', 'title': title, 'page': int(m.group(1))}

    return None


def _parse_toc_block(block: dict) -> list[dict]:
    """Parse a block into a list of TOC entry dicts."""
    entries = []
    btype = block.get('type', '')
    text  = (block.get('text') or '').strip()

    if btype == 'heading':
        if _PERIOD_RE.match(text):
            entries.append({'level': 'period', 'title': text, 'page': None})
        elif _WEEK_RE.match(text):
            entries.append({'level': 'week', 'title': text, 'page': None})
    else:
        for line in text.split('\n'):
            entry = _parse_toc_line(line)
            if entry:
                # A line that starts with a week heading is a sub-section
                clean = _BULLET_STRIP.sub('', line.strip())
                if _WEEK_RE.match(clean) and not _TAB_PAGE_RE.search(clean):
                    entry['level'] = 'week'
                entries.append(entry)

    return entries


def _inject_toc_table(pages: list[dict]) -> list[dict]:
    """Find the SOMMAIRE section and replace its blocks with a toc_table block."""
    # Locate SOMMAIRE heading
    sommaire_pi = sommaire_bi = None
    for pi, page in enumerate(pages):
        for bi, block in enumerate(page['blocks']):
            if (block.get('type') == 'heading' and
                    'SOMMAIRE' in (block.get('text') or '').upper()):
                sommaire_pi, sommaire_bi = pi, bi
                break
        if sommaire_pi is not None:
            break

    if sommaire_pi is None:
        return pages

    # Collect TOC blocks
    toc_entries: list[dict] = []
    toc_positions: set[tuple] = set()
    consecutive_non_toc = 0
    done = False

    for pi in range(sommaire_pi, len(pages)):
        if done:
            break
        start_bi = (sommaire_bi + 1) if pi == sommaire_pi else 0
        for bi in range(start_bi, len(pages[pi]['blocks'])):
            block = pages[pi]['blocks'][bi]
            text  = (block.get('text') or '').strip()
            btype = block.get('type', '')

            has_toc     = bool(re.search(r'\t\d+', text))
            is_section  = btype == 'heading' and (
                _PERIOD_RE.match(text) or _WEEK_RE.match(text)
            )

            if has_toc or is_section:
                parsed = _parse_toc_block(block)
                if parsed:
                    toc_entries.extend(parsed)
                    toc_positions.add((pi, bi))
                consecutive_non_toc = 0
            else:
                consecutive_non_toc += 1
                if consecutive_non_toc > 3 and toc_entries:
                    done = True
                    break

    if not toc_entries:
        return pages

    # Rebuild pages: replace TOC blocks with a single toc_table block
    new_pages = []
    inserted = False
    for pi, page in enumerate(pages):
        new_blocks = []
        for bi, block in enumerate(page['blocks']):
            if (pi, bi) in toc_positions:
                if not inserted:
                    new_blocks.append({'type': 'toc_table', 'entries': toc_entries})
                    inserted = True
                # skip original TOC block
            else:
                new_blocks.append(block)
        new_page = dict(page)
        new_page['blocks'] = new_blocks
        new_pages.append(new_page)

    return new_pages


# ── Block rendering ───────────────────────────────────────────────────────────

_SEMAINE_RE = re.compile(r'^((?:Semaine|Semame|Week)\s+[^\n]+)\n([\s\S]+)', re.IGNORECASE)


def _render_block(doc: Document, block: dict, lang: str, prev_was_bullet: bool = False) -> bool:
    """Render one block. Returns True if this block was a bullet paragraph."""
    btype = block.get("type")

    if btype in ("qr_label", "caption"):
        return False

    if btype == "image_table":
        add_image_table(doc, block.get("descriptions", []), lang=lang)
        return False

    if btype == "obj_mat_table":
        add_obj_mat_table(
            doc,
            obj_label=block.get("obj_label", ""),
            obj_blocks=block.get("obj_blocks", []),
            mat_label=block.get("mat_label", ""),
            mat_blocks=block.get("mat_blocks", []),
            lang=lang,
        )
        return False

    if btype == "materiel_table":
        add_materiel_table(
            doc,
            ps_ressources=block.get("ps_ressources", []),
            ms_ressources=block.get("ms_ressources", []),
            ps_classe=block.get("ps_classe", []),
            ms_classe=block.get("ms_classe", []),
            ressources_blocks=block.get("ressources_blocks", []),
            classe_blocks=block.get("classe_blocks", []),
            lang=lang,
        )
        return False

    if btype == "section_box":
        from config import Colors
        border_color = Colors.BLUE if block.get("border_color") == "blue" else Colors.ORANGE
        add_section_box(
            doc,
            label=block.get("label", ""),
            content_blocks=block.get("content_blocks", []),
            border_color=border_color,
            lang=lang,
        )
        return False

    if btype == "prog_banner":
        add_prog_banner(
            doc,
            periode_num=block.get("periode_num", 1),
            label=block.get("label", "Programmation"),
            classes=block.get("classes", ["PS", "MS"]),
        )
        return False

    if btype == "activity_title":
        week          = block.get("week") or ""
        periode_num   = block.get("periode_num")
        activity_type = block.get("activity_type", "")
        activity_label = block.get("activity_label", "")
        label_lower   = activity_label.lower()

        if (week and periode_num and re.search(r'\d', str(week)) and
                _is_main_activity_label(activity_label) and
                activity_type in ("ritualisee", "guidee", "autonomie")):
            # Standard activity banner
            add_activity_title_bar(
                doc,
                periode_num=int(periode_num),
                week=week,
                activity_label=activity_label,
                activity_type=activity_type,
                classes=block.get("classes", []),
            )
        elif (week and periode_num and re.search(r'\d', str(week)) and
              ('programmation' in label_lower or 'organisation' in label_lower)):
            # Programmation / Organisation des ateliers banner — uses period color
            add_prog_banner(
                doc,
                periode_num=int(periode_num),
                label=activity_label,
                week=week,
                classes=block.get("classes", ["PS", "MS"]),
            )
        elif activity_label:
            add_heading(doc, activity_label, level=2, lang=lang)
        return False

    if btype == "subtitle_badge":
        add_subtitle_badge(
            doc,
            badge_text=block.get("badge_text", ""),
            title_text=block.get("title_text", ""),
            activity_type=block.get("activity_type", "ritualisee"),
            lang=lang,
        )
        return False

    if btype == "image_placeholder":
        desc = block.get("description", "").strip()
        add_image_placeholder(doc, f"[ IMAGE — {desc} ]" if desc else "[ IMAGE ]")
        return False

    if btype == "toc_table":
        add_toc_table(doc, block.get("entries", []))
        return False

    if btype == "heading":
        level = block.get("level", 1)
        raw   = block.get("text") or ""
        text  = str(raw) if not isinstance(raw, str) else raw
        if _is_period_line(text):
            num = _extract_period_num(text)
            add_period_banner(doc, num)
        elif _is_savoir_box(text):
            add_info_box(doc, text, lang)
        elif _is_section_label(text):
            add_section_label(doc, text, lang)
        else:
            add_heading(doc, text, level, lang,
                        bg_color=block.get("bg_color"),
                        text_color=block.get("text_color"))
        return False

    elif btype == "paragraph":
        raw  = block.get("text") or ""
        text = (str(raw) if not isinstance(raw, str) else raw).strip()
        if not text:
            return prev_was_bullet

        if _is_page_number(text):
            return prev_was_bullet

        if _is_url_only(text):
            return prev_was_bullet

        # Skip image/QR captions prefixed with ► or ▸
        if text.startswith(('►', '▸', '► ', '▸ ')):
            return prev_was_bullet

        text = _strip_toc_refs(text)
        text = _strip_page_refs(text)
        if not text:
            return prev_was_bullet

        m = _SEMAINE_RE.match(text)
        if m:
            add_heading(doc, m.group(1).strip(), level=2, lang=lang)
            rest = m.group(2).strip()
            if rest:
                add_paragraph(doc, rest, lang)
            return _text_is_bullet(rest)

        is_bullet = _text_is_bullet(text)
        if not is_bullet and prev_was_bullet:
            add_indented_paragraph(doc, text, lang)
        else:
            add_paragraph(doc, text, lang)
        return is_bullet

    elif btype == "table":
        data = block.get("data", [])
        add_table(doc, data, col_colors=block.get("col_colors"),
                  compact=block.get("compact", False))
        return False

    elif btype == "image":
        return False

    return False


# ── Helpers ───────────────────────────────────────────────────────────────────

def _mark_qr_labels(blocks: list[dict]) -> list[dict]:
    """Tag short non-bullet paragraphs that sit directly before a URL block as qr_label."""
    result = []
    n = len(blocks)
    for i, block in enumerate(blocks):
        if (block.get('type') == 'paragraph' and i + 1 < n):
            next_raw  = blocks[i + 1].get('text') or ''
            next_text = str(next_raw).strip() if not isinstance(next_raw, str) else next_raw.strip()
            if _is_url_only(next_text):
                text_raw = block.get('text') or ''
                text = str(text_raw).strip() if not isinstance(text_raw, str) else text_raw.strip()
                if not text.startswith(('•', '-')) and len(text.split()) <= 10:
                    result.append({**block, 'type': 'qr_label'})
                    continue
        result.append(block)
    return result


def _text_is_bullet(text: str) -> bool:
    first = next((l.strip() for l in text.splitlines() if l.strip()), "")
    return first.startswith("•") or bool(re.match(r'^-\s', first))


def _is_page_number(text: str) -> bool:
    return bool(re.fullmatch(r'\d{1,3}', text.strip()))


_URL_RE = re.compile(
    r'^(https?://|www\.|\w[\w.-]+\.[a-z]{2,}(/|\s|$))',
    re.IGNORECASE | re.MULTILINE,
)

def _is_url_only(text: str) -> bool:
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if not lines:
        return False
    return all(_URL_RE.match(l) for l in lines)


_TOC_REF_RE = re.compile(r'\t\d+\.?\s*$', re.MULTILINE)

def _strip_toc_refs(text: str) -> str:
    return _TOC_REF_RE.sub('', text).strip()


_PAGE_REF_RE = re.compile(
    r'\(?\bp\.?\s*\d+(?:\s*[àa]\s*\d+)?\)?'   # p. 92, (p. 136 à 139)
    r'|,?\s*voir\s+page\s+\w+'                  # voir page X
    r'|,?\s*see\s+page\s+\w+',                  # see page X
    re.IGNORECASE,
)

def _strip_page_refs(text: str) -> str:
    """Remove inline page cross-references (p. XX, voir page...) from body text."""
    cleaned = _PAGE_REF_RE.sub('', text)
    return re.sub(r'[ \t]{2,}', ' ', cleaned).strip()


_MAIN_ACTIVITY_LABELS = {
    "activités ritualisées", "activité ritualisée",
    "apprentissages guidés", "apprentissages guides",
    "autonomie et semi-autonomie",
}

def _is_main_activity_label(label: str) -> bool:
    return label.strip().lower() in _MAIN_ACTIVITY_LABELS

_SECTION_LABELS = (
    'objectif', 'déroulement', 'deroulement', 'différenciation', 'differentiation',
    'matériel', 'materiel', 'organisation', 'ressource', 'remarque',
)

def _is_section_label(text: str) -> bool:
    clean = text.strip().rstrip(':').lower()
    return any(clean.startswith(s) for s in _SECTION_LABELS)

def _is_savoir_box(text: str) -> bool:
    t = text.lower().replace('’', "'").replace('‘', "'")
    return "ce qu'il faut savoir" in t

def _is_period_line(text: str) -> bool:
    return bool(re.match(r'^(PÉRIODE|PERIOD)\s*\d+', text.strip(), re.IGNORECASE))


def _extract_period_num(text: str) -> int:
    m = re.search(r"\d+", text)
    return int(m.group()) if m else 0
