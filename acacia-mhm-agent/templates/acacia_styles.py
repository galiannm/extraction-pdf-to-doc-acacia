"""
Helpers to apply Acacia brand styles to a python-docx Document.
"""
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_PERIOD_ENTRY_RE = re.compile(r'^(PÉRIODE|PERIOD)\s*\d+', re.IGNORECASE)
_WEEK_ENTRY_RE   = re.compile(r'^(Semaine|Semame|Week)\s+', re.IGNORECASE)

from config import Colors, Fonts, FaithfulFonts, LOGO_PATH, PERIOD_COLORS, ACTIVITY_COLORS

# ── Mode toggle (faithful vs Acacia-branded) ──────────────────────────────────
_fonts: type = Fonts
_branded: bool = True

def _set_mode(faithful: bool) -> None:
    global _fonts, _branded
    _fonts = FaithfulFonts if faithful else Fonts
    _branded = not faithful

# ── Page layout ──────────────────────────────────────────────────────────────

PAGE_WIDTH   = Inches(8.27)   # A4
PAGE_HEIGHT  = Inches(11.69)
MARGIN       = Inches(1.0)
CONTENT_WIDTH = Inches(6.27)  # PAGE_WIDTH - 2*MARGIN


def setup_document(doc: Document, lang: str = 'fr-FR') -> None:
    """Apply page size, margins, base paragraph font, and proofing language."""
    section = doc.sections[0]
    section.page_width  = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, MARGIN)

    normal = doc.styles["Normal"]
    normal.font.name = _fonts.BODY
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(*Colors.DARK_GREY)
    _set_style_lang(normal, lang)
    _set_doc_default_lang(doc, lang)


# ── Header / Footer ───────────────────────────────────────────────────────────

def add_header(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    header = section.header
    for p in header.paragraphs:
        p.clear()
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if _branded and LOGO_PATH.exists():
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), height=Pt(30))


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    for p in footer.paragraphs:
        p.clear()
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if _branded:
        run = p.add_run("www.acacia-education.com    ")
        run.font.name = _fonts.BODY
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*Colors.DARK_GREY)

    _add_page_number_field(p, size=9)


def _add_page_number_field(para, size: int = 9) -> None:
    """Append a Word PAGE field run to an existing paragraph."""
    run = para.add_run()
    run.font.name = _fonts.BODY
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*Colors.DARK_GREY)
    r = run._r
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    r.append(begin)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    r.append(instr)
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    r.append(end)


# ── Cover page ────────────────────────────────────────────────────────────────

def add_cover_page(doc: Document, title: str, subtitle: str = "") -> None:
    """Cover page: Acacia-branded (with logo) or plain depending on mode."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    if _branded and LOGO_PATH.exists():
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(2.5))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    run = p.add_run(title)
    run.font.name = _fonts.HEADING
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*Colors.ORANGE)

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(subtitle)
        run.font.name = _fonts.BODY
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(*Colors.DARK_GREY)

    doc.add_page_break()


# ── Period cover banner ───────────────────────────────────────────────────────

def add_period_banner(doc: Document, period_num: int, date_range: str = "") -> None:
    """Coloured banner row to mark a new period."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _shade_cell(cell, _rgb_hex(Colors.ORANGE))
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(f"Period {period_num}")
    run.font.name  = _fonts.HEADING
    run.font.size  = Pt(28)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(*Colors.WHITE)
    if date_range:
        run2 = p.add_run(f"  {date_range}")
        run2.font.name  = _fonts.BODY
        run2.font.size  = Pt(14)
        run2.font.color.rgb = RGBColor(*Colors.WHITE)
    _remove_table_borders(table)
    doc.add_paragraph()


# ── MHM activity title bar ────────────────────────────────────────────────────

def _normalize_week(week: str) -> str:
    """Normalize 'SEMAINE 1' / 'Semaine 1' → 'S1', leave 'S1' unchanged."""
    import re as _re
    m = _re.search(r'\d+', str(week))
    if m and len(str(week)) > 3:
        return f"S{m.group()}"
    return str(week)


def add_activity_title_bar(
    doc: Document,
    periode_num: int,
    week: str,
    activity_label: str,
    activity_type: str,
    classes: list[str],
) -> None:
    """Render the composite MHM title banner: [Période N / S[week]] [Activity type] [PS] [MS]"""
    p_color = PERIOD_COLORS.get(periode_num, Colors.BLUE)
    a_color = ACTIVITY_COLORS.get(activity_type, Colors.YELLOW)
    week = _normalize_week(week)

    badges = [c for c in classes if c in ("PS", "MS")]
    n_cols = 2 + len(badges)

    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"

    content_tw = int(CONTENT_WIDTH * 1440 / 914400)
    period_w   = int(0.9 * 1440)
    badge_w    = int(0.55 * 1440)
    activity_w = content_tw - period_w - badge_w * len(badges)
    _set_table_col_widths(table, [period_w, activity_w] + [badge_w] * len(badges))

    # Period / week cell — white background, period-colored text
    cell = table.cell(0, 0)
    _shade_cell(cell, "FFFFFF")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(f"Période {periode_num}\n")
    r1.font.name  = _fonts.BODY
    r1.font.size  = Pt(7)
    r1.font.bold  = True
    r1.font.color.rgb = RGBColor(*p_color)
    _set_run_lang(r1, 'fr-FR')
    r2 = p.add_run(week)
    r2.font.name  = _fonts.HEADING
    r2.font.size  = Pt(20)
    r2.font.bold  = True
    r2.font.color.rgb = RGBColor(*p_color)
    _set_run_lang(r2, 'fr-FR')

    # Activity type cell — activity color background, white text
    cell = table.cell(0, 1)
    _shade_cell(cell, _rgb_hex(a_color))
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Pt(10)
    run = p.add_run(activity_label)
    run.font.name  = _fonts.HEADING
    run.font.size  = Pt(20)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(*Colors.WHITE)
    _set_run_lang(run, 'fr-FR')

    # Class badge cells
    badge_colors = {"PS": Colors.PS_BADGE, "MS": Colors.MS_BADGE}
    for i, cls in enumerate(badges):
        cell = table.cell(0, 2 + i)
        _shade_cell(cell, _rgb_hex(badge_colors[cls]))
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(cls)
        run.font.name  = _fonts.HEADING
        run.font.size  = Pt(14)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(*Colors.WHITE)
        _set_run_lang(run, 'fr-FR')

    _remove_table_borders(table)
    doc.add_paragraph()


def add_subtitle_badge(
    doc: Document,
    badge_text: str,
    title_text: str,
    activity_type: str,
    lang: str = 'fr-FR',
) -> None:
    """Render a day/timing badge + subtitle on one line: [CHAQUE JOUR] Title text"""
    a_color = ACTIVITY_COLORS.get(activity_type, Colors.YELLOW)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    content_tw = int(CONTENT_WIDTH * 1440 / 914400)
    badge_w    = int(1.4 * 1440)
    title_w    = content_tw - badge_w
    _set_table_col_widths(table, [badge_w, title_w])

    # Badge cell
    cell = table.cell(0, 0)
    _shade_cell(cell, _rgb_hex(a_color))
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(badge_text)
    run.font.name  = _fonts.BODY
    run.font.size  = Pt(10)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(*Colors.WHITE)
    _set_run_lang(run, lang)

    # Title text cell
    cell = table.cell(0, 1)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(8)
    run = p.add_run(title_text)
    run.font.name  = _fonts.HEADING
    run.font.size  = Pt(14)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(*a_color)
    _set_run_lang(run, lang)

    _remove_table_borders(table)


# ── Language helpers ─────────────────────────────────────────────────────────

def _set_style_lang(style, lang: str) -> None:
    rPr = style.element.get_or_add_rPr()
    lang_el = rPr.find(qn('w:lang'))
    if lang_el is None:
        lang_el = OxmlElement('w:lang')
        rPr.append(lang_el)
    lang_el.set(qn('w:val'), lang)


def _set_doc_default_lang(doc: Document, lang: str) -> None:
    styles_el = doc.styles.element
    docDefaults = styles_el.find(qn('w:docDefaults'))
    if docDefaults is None:
        return
    rPrDefault = docDefaults.find(qn('w:rPrDefault'))
    if rPrDefault is None:
        rPrDefault = OxmlElement('w:rPrDefault')
        docDefaults.append(rPrDefault)
    rPr = rPrDefault.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        rPrDefault.append(rPr)
    lang_el = rPr.find(qn('w:lang'))
    if lang_el is None:
        lang_el = OxmlElement('w:lang')
        rPr.append(lang_el)
    lang_el.set(qn('w:val'), lang)


def _set_run_lang(run, lang: str) -> None:
    rPr = run._r.get_or_add_rPr()
    lang_el = rPr.find(qn('w:lang'))
    if lang_el is None:
        lang_el = OxmlElement('w:lang')
        rPr.append(lang_el)
    lang_el.set(qn('w:val'), lang)


# ── Term highlighting ─────────────────────────────────────────────────────────

# Each entry: (regex_pattern, color_rgb, bold)
_HIGHLIGHT_TERMS = [
    # Activity types — FR
    (r'activités?\s+ritualisées?',           Colors.RITUALISED_ACTIVITIES, True),
    (r'activités?\s+guidées?',               Colors.GUIDED_LEARNING,       True),
    (r'activités?\s+semi-autonomes?',        Colors.AUTONOMOUS_ACTIVITIES, True),
    (r'activités?\s+autonomes?',             Colors.AUTONOMOUS_ACTIVITIES, True),
    # Activity types — EN
    (r'ritualiz(?:ed)?\s+activit(?:y|ies)',  Colors.RITUALISED_ACTIVITIES, True),
    (r'guided\s+(?:learning|activit\w*)',    Colors.GUIDED_LEARNING,       True),
    (r'semi-autonomous\s+activit\w*',        Colors.AUTONOMOUS_ACTIVITIES, True),
    (r'autonomous\s+activit\w*',             Colors.AUTONOMOUS_ACTIVITIES, True),
    # Assessment — FR & EN
    (r'évaluation\s+formative',              Colors.ORANGE, True),
    (r'évaluation\s+sommative',              Colors.ORANGE, True),
    (r'formative\s+assessment',              Colors.ORANGE, True),
    (r'summative\s+assessment',              Colors.ORANGE, True),
    # Grade labels — FR
    (r'Petite\s+Section',                    Colors.PS_BADGE,  True),
    (r'Moyenne\s+Section',                   Colors.MS_BADGE,  True),
    # Method name
    (r'\bMHM\b',                             Colors.ORANGE,    True),
    (r'\bheuristique\b',                     Colors.ORANGE,    True),
]

# Compile once
_HIGHLIGHT_RE = re.compile(
    '(' + '|'.join(p for p, _, _ in _HIGHLIGHT_TERMS) + ')',
    re.IGNORECASE,
)
_TERM_COLOR: dict[str, tuple] = {}  # cache lowercase → (color, bold)
_BOLD_RE = re.compile(r'\*\*(.+?)\*\*', re.DOTALL)


def _highlight_lookup(matched: str) -> tuple[tuple, bool]:
    key = matched.lower()
    if key not in _TERM_COLOR:
        for pattern, color, bold in _HIGHLIGHT_TERMS:
            if re.fullmatch(pattern, matched, re.IGNORECASE):
                _TERM_COLOR[key] = (color, bold)
                break
        else:
            _TERM_COLOR[key] = (None, False)
    return _TERM_COLOR[key]


def _parse_bold_segments(text: str) -> list[tuple[str, bool]]:
    """Split text on **...** markers → list of (segment, is_bold)."""
    segments: list[tuple[str, bool]] = []
    last = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > last:
            segments.append((text[last:m.start()], False))
        segments.append((m.group(1), True))
        last = m.end()
    if last < len(text):
        segments.append((text[last:], False))
    return segments or [(text, False)]


def _sanitize_text(text: str) -> str:
    """Strip null bytes and XML-incompatible control characters, keep \\n and \\t."""
    return ''.join(ch for ch in text if ch == '\n' or ch == '\t' or (ord(ch) >= 32 and ch != '\x7f'))


def _add_highlighted_runs(p, text: str, base_size: int, lang: str) -> None:
    """Add runs to paragraph `p`, respecting **bold** markup and MHM term colours."""
    text = _sanitize_text(text)
    for seg_text, is_markup_bold in _parse_bold_segments(text):
        for chunk in _HIGHLIGHT_RE.split(seg_text):
            if not chunk:
                continue
            run = p.add_run(chunk)
            run.font.name = _fonts.BODY
            run.font.size = Pt(base_size)
            color, _ = _highlight_lookup(chunk)
            if color:
                run.font.color.rgb = RGBColor(*color)
            if is_markup_bold:
                run.font.bold = True
            _set_run_lang(run, lang)


# ── Text helpers ──────────────────────────────────────────────────────────────

def _fix_soft_hyphens(text: str) -> str:
    """Join words split by PDF line-break hyphens: 'manipu- lations' or 'manipu-\nlations' → 'manipulations'."""
    # Match hyphen + optional spaces/newline + lowercase letter (not a real hyphen like "enseignant-e")
    return re.sub(r'-[\s\n]+([a-zàâäéèêëîïôùûüçœæ])', r'\1', text)


def _is_bullet_line(line: str) -> bool:
    s = line.strip()
    return s.startswith('•') or bool(re.match(r'^-\s', s))


# ── Text blocks ───────────────────────────────────────────────────────────────

def add_section_label(doc: Document, text: str, lang: str = 'fr-FR') -> None:
    """Render Objectif / Déroulement / Différenciation etc. as bold + underline black."""
    clean = re.sub(r'\*\*(.+?)\*\*', r'\1', _fix_soft_hyphens(text)).strip()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(clean)
    run.font.name      = _fonts.BODY
    run.font.size      = Pt(11)
    run.font.bold      = True
    run.font.underline = True
    run.font.color.rgb = RGBColor(*Colors.DARK_GREY)
    _set_run_lang(run, lang)


def add_info_box(doc: Document, text: str, lang: str = 'fr-FR') -> None:
    """Render 'Ce qu'il faut savoir' as a white box with blue border and blue text."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    content_tw = int(CONTENT_WIDTH * 1440 / 914400)
    box_w = int(content_tw * 0.55)   # ~55% width — compact, not full-width
    _set_table_col_widths(table, [box_w])
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(6)
    run = p.add_run(text.strip())
    run.font.name  = _fonts.BODY
    run.font.size  = Pt(10)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(*Colors.BLUE)
    _set_run_lang(run, lang)
    # Apply blue border to the table
    _set_table_border_color(table, _rgb_hex(Colors.BLUE))
    doc.add_paragraph()


def _set_table_border_color(table, hex_color: str) -> None:
    """Set all table borders to a specific color."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:color'), hex_color)
        tblBorders.append(border)
    tblPr.append(tblBorders)


def add_heading(doc: Document, text: str, level: int, lang: str = 'fr-FR',
                bg_color: str | None = None, text_color: str | None = None) -> None:
    sizes  = {1: 19, 2: 16, 3: 13}
    colors = {1: Colors.ORANGE, 2: Colors.BLUE, 3: Colors.DARK_GREY}
    # Strip **bold** markers — headings are already bold
    clean = re.sub(r'\*\*(.+?)\*\*', r'\1', _fix_soft_hyphens(text)).replace('\n', ' ').strip()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(clean)
    run.font.name  = _fonts.HEADING
    run.font.size  = Pt(sizes.get(level, 13))
    run.font.bold  = True

    if not _branded and bg_color:
        hex_bg = _normalize_color(bg_color)
        if hex_bg and hex_bg.upper() not in ("FFFFFF", "F0F0F0", "EEEEEE"):
            # Use the PDF's color as text color — never shade the full paragraph width
            run.font.color.rgb = RGBColor(*tuple(int(hex_bg[i:i+2], 16) for i in (0, 2, 4)))
        else:
            run.font.color.rgb = RGBColor(*colors.get(level, Colors.DARK_GREY))
    else:
        run.font.color.rgb = RGBColor(*colors.get(level, Colors.DARK_GREY))

    _set_run_lang(run, lang)


def add_paragraph(doc: Document, text: str, lang: str = 'fr-FR') -> None:
    text = _fix_soft_hyphens(text)
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if not lines:
        return

    has_bullets = any(_is_bullet_line(l) for l in lines)

    if not has_bullets:
        _add_body_para(doc, ' '.join(lines), lang)
        return

    # Group lines: each bullet + continuations → one bullet chunk;
    # non-bullet runs → body chunk.
    chunks: list[tuple[bool, str]] = []
    current_bullet = None
    current_lines: list[str] = []

    for line in lines:
        if _is_bullet_line(line):
            if current_lines:
                chunks.append((bool(current_bullet), ' '.join(current_lines)))
            current_bullet = True
            current_lines = [re.sub(r'^[•\-]\s*', '', line)]
        else:
            if current_bullet is True:
                current_lines.append(line)
            else:
                if current_bullet is False:
                    current_lines.append(line)
                else:
                    current_bullet = False
                    current_lines = [line]

    if current_lines:
        chunks.append((bool(current_bullet), ' '.join(current_lines)))

    for is_bullet, chunk_text in chunks:
        if is_bullet:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.first_line_indent = Pt(-12)
            p.paragraph_format.space_after = Pt(3)
            _add_highlighted_runs(p, '•  ' + chunk_text, 11, lang)
        else:
            _add_body_para(doc, chunk_text, lang)


def _add_body_para(doc: Document, text: str, lang: str = 'fr-FR', left_indent: int = 0) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if left_indent:
        p.paragraph_format.left_indent = Pt(left_indent)
    _add_highlighted_runs(p, text, 11, lang)


def add_indented_paragraph(doc: Document, text: str, lang: str = 'fr-FR') -> None:
    """Body paragraph indented to sit under a preceding bullet item."""
    text = _fix_soft_hyphens(text)
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if lines:
        _add_body_para(doc, ' '.join(lines), lang, left_indent=18)


# ── TOC table ─────────────────────────────────────────────────────────────────

def add_toc_table(doc: Document, entries: list[dict]) -> None:
    """Render sommaire as a 2-column table: Title | Page."""
    if not entries:
        return

    table = doc.add_table(rows=len(entries), cols=2)
    table.style = "Table Grid"

    content_twips = int(CONTENT_WIDTH * 1440 / 914400)
    page_col_w  = int(0.55 * 1440)   # ~0.55" for page numbers
    title_col_w = content_twips - page_col_w
    _set_table_col_widths(table, [title_col_w, page_col_w])

    for r_idx, entry in enumerate(entries):
        level    = entry.get('level', 'item')
        title    = entry.get('title', '')
        page_num = entry.get('page')

        t_cell = table.cell(r_idx, 0)
        p_cell = table.cell(r_idx, 1)
        t_cell.text = ""
        p_cell.text = ""

        tp = t_cell.paragraphs[0]
        pp = p_cell.paragraphs[0]

        for para in (tp, pp):
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(3)

        pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if level == 'period':
            _shade_cell(t_cell, _rgb_hex(Colors.ORANGE))
            _shade_cell(p_cell, _rgb_hex(Colors.ORANGE))
            run = tp.add_run(title.upper())
            run.font.name  = _fonts.HEADING
            run.font.size  = Pt(11)
            run.font.bold  = True
            run.font.color.rgb = RGBColor(*Colors.WHITE)

        elif level == 'week':
            _shade_cell(t_cell, _rgb_hex(Colors.BLUE))
            _shade_cell(p_cell, _rgb_hex(Colors.BLUE))
            tp.paragraph_format.left_indent = Pt(14)
            run = tp.add_run(title)
            run.font.name  = _fonts.HEADING
            run.font.size  = Pt(10)
            run.font.bold  = True
            run.font.color.rgb = RGBColor(*Colors.WHITE)

        else:  # item
            tp.paragraph_format.left_indent = Pt(28)
            run = tp.add_run(title)
            run.font.name  = _fonts.BODY
            run.font.size  = Pt(10)
            run.font.color.rgb = RGBColor(*Colors.DARK_GREY)

            if page_num is not None:
                pr = pp.add_run(str(page_num))
                pr.font.name  = _fonts.BODY
                pr.font.size  = Pt(10)
                pr.font.color.rgb = RGBColor(*Colors.DARK_GREY)

    doc.add_paragraph()


# ── Tables ─────────────────────────────────────────────────────────────────────

def add_table(doc: Document, data: list[list[str]], col_colors: list[str] | None = None) -> None:
    """Render an extracted table with colour coding and proper column widths."""
    if not data:
        return

    num_cols = max(len(row) for row in data)
    if num_cols == 0:
        return

    table = doc.add_table(rows=len(data), cols=num_cols)
    table.style = "Table Grid"

    content_twips = int(CONTENT_WIDTH * 1440 / 914400)
    _set_table_col_widths(table, _proportional_col_widths(data, num_cols, content_twips))

    for r_idx, row in enumerate(data):
        for c_idx in range(num_cols):
            cell_text = _safe_text(row[c_idx]) if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

            # Render cell text — handle embedded bullets, newlines, and **bold** markers
            cell_lines = _fix_soft_hyphens(cell_text).split('\n') if cell_text else ['']
            first = True
            for line in cell_lines:
                line = line.strip()
                if not line:
                    continue
                if not first:
                    p.add_run().add_break()
                for seg, is_bold in _parse_bold_segments(line):
                    if not seg:
                        continue
                    run = p.add_run(seg)
                    run.font.name = _fonts.BODY
                    run.font.size = Pt(9)
                    if is_bold:
                        run.font.bold = True
                first = False

            # Faithful mode: use captured col_colors for header row
            if not _branded and col_colors and r_idx == 0 and c_idx < len(col_colors):
                hex_fill = _normalize_color(col_colors[c_idx])
                if hex_fill:
                    _shade_cell(cell, hex_fill)
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(*Colors.WHITE)
                        run.font.bold = True
            else:
                fill = _cell_fill(r_idx, c_idx, cell_text, data)
                if fill:
                    _shade_cell(cell, fill)
                    if r_idx == 0 or _is_week_row(cell_text):
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(*Colors.WHITE)
                            run.font.bold = True

    doc.add_paragraph()


def _proportional_col_widths(data: list, num_cols: int, content_twips: int) -> list[int]:
    """Distribute column widths proportional to max cell content length per column."""
    col_len = [0] * num_cols
    for row in data:
        for c, cell in enumerate(row):
            if c < num_cols:
                col_len[c] = max(col_len[c], len(str(cell or '')))

    # Clamp each column: min 4 chars, max 120 chars equivalent
    col_len = [max(4, min(120, l)) for l in col_len]
    total = sum(col_len)

    # Minimum column: 6% of content width
    min_twips = max(int(content_twips * 0.06), 600)
    raw = [int(content_twips * l / total) for l in col_len]

    # Enforce minimums, redistribute surplus
    result = [max(r, min_twips) for r in raw]
    diff = content_twips - sum(result)
    # Distribute rounding diff to the widest column
    if result:
        result[result.index(max(result))] += diff
    return result


def _set_table_col_widths(table, twips: list) -> None:
    """Set explicit column widths (in twips) via OOXML for reliable Word rendering."""
    tbl = table._tbl

    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(sum(twips)))
    tblW.set(qn('w:type'), 'dxa')

    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.insert(1, tblGrid)
    else:
        for child in list(tblGrid):
            tblGrid.remove(child)
    for w in twips:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(w))
        tblGrid.append(gridCol)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(twips):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                tcW.set(qn('w:w'), str(twips[i]))
                tcW.set(qn('w:type'), 'dxa')


def _safe_text(value) -> str:
    """Coerce any LLM cell value to a plain string."""
    if value is None:
        return ""
    if isinstance(value, list):
        return "\n".join(_safe_text(v) for v in value)
    if isinstance(value, dict):
        return str(value.get("text", value))
    return str(value)


def _cell_fill(r_idx: int, c_idx: int, text, data: list) -> str | None:
    text = _safe_text(text)
    t = text.lower().strip()

    if _is_week_row(text):
        return _rgb_hex(Colors.WEEK_ROW)

    if r_idx == 0:
        if any(k in t for k in ("ritualis", "ritualized", "ritual")):
            return _rgb_hex(Colors.RITUALISED_ACTIVITIES)
        if any(k in t for k in ("apprentiss", "guided", "learning")):
            return _rgb_hex(Colors.GUIDED_LEARNING)
        if any(k in t for k in ("autonomie", "autonomous", "semi")):
            return _rgb_hex(Colors.AUTONOMOUS_ACTIVITIES)
        if any(k in t for k in ("objectif", "objective")):
            return _rgb_hex(Colors.OBJECTIVES_HEADER)
        if t in ("ps", "ms"):
            return _rgb_hex(Colors.PS_BADGE if t == "ps" else Colors.MS_BADGE)
        return _rgb_hex(Colors.PURPLE)

    if t == "ps":
        return _rgb_hex(Colors.PS_BADGE)
    if t == "ms":
        return _rgb_hex(Colors.MS_BADGE)

    if data and c_idx < len(data[0]):
        header = _safe_text(data[0][c_idx]).lower()
        if any(k in header for k in ("ritualis", "ritualized")):
            return _rgb_hex(Colors.LIGHT_GREY)
        if any(k in header for k in ("apprentiss", "guided")):
            return _rgb_hex(Colors.LIGHT_GREY)

    return None


def _is_week_row(text: str) -> bool:
    t = text.strip().upper()
    return t.startswith("SEMAINE") or t.startswith("WEEK")


# ── Images ────────────────────────────────────────────────────────────────────

def add_image_placeholder(doc: Document, text: str) -> None:
    """Render an image placeholder as a light-grey shaded box with italic text."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _shade_cell(cell, _rgb_hex(Colors.LIGHT_GREY))
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(text)
    run.font.name   = _fonts.BODY
    run.font.size   = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(*Colors.DARK_GREY)
    doc.add_paragraph()


def add_image(doc: Document, abs_path: str, caption: str = "") -> None:
    from pathlib import Path
    img = Path(abs_path)
    if not img.exists():
        doc.add_paragraph(f"[Image not found: {img.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture(str(img), width=Inches(4))
    except Exception:
        p.add_run(f"[Could not embed: {img.name}]")
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cp.runs:
            r.font.size = Pt(9)
            r.font.italic = True


# ── Color helpers ─────────────────────────────────────────────────────────────

_COLOR_NAMES = {
    "yellow":  "F5C842",
    "blue":    "5BC8F5",
    "green":   "8DC63F",
    "orange":  "F5A623",
    "purple":  "B07CC6",
    "red":     "E8786A",
    "coral":   "E8786A",
    "grey":    "AAAAAA",
    "gray":    "AAAAAA",
    "white":   "FFFFFF",
    "black":   "404040",
    "none":    None,
}

def _normalize_color(color: str) -> str | None:
    """Convert a color name or #hex string to a 6-char hex string, or None."""
    if not color:
        return None
    c = color.strip().lower()
    if c in _COLOR_NAMES:
        return _COLOR_NAMES[c]
    if c.startswith('#') and len(c) in (4, 7):
        return c[1:].upper()
    return None


def _shade_paragraph(p, hex_color: str) -> None:
    """Apply a background fill to a paragraph via OOXML."""
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)


# ── XML helpers ───────────────────────────────────────────────────────────────

def _rgb_hex(rgb: tuple) -> str:
    return "{:02X}{:02X}{:02X}".format(*rgb)


def _shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _remove_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)
    tblPr.append(tblBorders)
